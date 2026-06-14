"""
文章分析报告 Word 文档生成脚本（python-docx）

用法：
  echo '<JSON>' | python generate_report.py
  python generate_report.py --input data.json
  python generate_report.py --input data.json --output report.docx

JSON 输入格式见 SKILL.md

依赖：pip install python-docx
"""

import sys
import json
import argparse

from docx import Document
from docx.shared import Pt, Cm, RGBColor, DXA
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ============================================================
# 常量
# ============================================================

HEADING1_COLOR = RGBColor(0x1F, 0x38, 0x64)  # 深蓝
HEADING2_COLOR = RGBColor(0x2E, 0x75, 0xB6)  # 蓝
HEADING3_COLOR = RGBColor(0x37, 0x56, 0x23)  # 深绿
BODY_FONT = 'Arial'
BODY_SIZE = Pt(11)


# ============================================================
# 读取输入
# ============================================================

def load_input():
    parser = argparse.ArgumentParser(description='Generate article analysis report')
    parser.add_argument('--input', help='JSON input file path')
    parser.add_argument('--output', help='Output .docx file path')
    args = parser.parse_args()

    if args.input:
        with open(args.input, 'r', encoding='utf-8') as f:
            data = json.load(f)
    else:
        if sys.stdin.isatty():
            print('用法: echo \'<JSON>\' | python generate_report.py', file=sys.stderr)
            print('      python generate_report.py --input data.json [--output report.docx]', file=sys.stderr)
            sys.exit(1)
        data = json.load(sys.stdin)

    data['outputFile'] = args.output or data.get('outputFile', 'article_analysis_report.docx')
    return data


# ============================================================
# 文档构建
# ============================================================

def set_run_font(run, size=BODY_SIZE, color=None, bold=False):
    """统一设置 run 的字体属性"""
    run.font.name = BODY_FONT
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置中文字体（兼容东亚文字）
    run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)


def add_paragraph(doc, text, size=BODY_SIZE, color=None, bold=False, alignment=None):
    """添加普通段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, size, color, bold)
    return p


def add_heading1(doc, text):
    """添加 Heading 1（深蓝）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, Pt(16), HEADING1_COLOR, bold=True)
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_heading2(doc, text):
    """添加 Heading 2（蓝）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, Pt(13), HEADING2_COLOR, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_heading3(doc, text):
    """添加 Heading 3（深绿）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, Pt(11), HEADING3_COLOR, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_spacer(doc):
    """添加空行"""
    p = doc.add_paragraph()
    run = p.add_run('')
    set_run_font(run, Pt(6))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_bullet(doc, text):
    """添加项目符号条目"""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_run_font(run, BODY_SIZE)
    return p


def add_table(doc, sec):
    """添加表格"""
    columns = sec.get('columns', [])
    rows = sec.get('rows', [])
    if not columns or not rows:
        return

    col_count = len(columns)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = [c if isinstance(c, str) else c.get('header', '') for c in columns]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        set_run_font(run, BODY_SIZE, bold=True)
        # 浅蓝背景
        shading = cell._element.get_or_add_tcPr()
        shd_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D5E8F0',
            qn('w:val'): 'clear'
        })
        shading.append(shd_elm)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            text = cell_text if isinstance(cell_text, str) else cell_text.get('text', '')
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(run, BODY_SIZE)

    add_spacer(doc)


def build_document(data):
    """构建完整文档"""
    doc = Document()

    # 页面设置：A4，边距约 2.1cm
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    # 页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(data.get('title', '文章分析报告'))
    set_run_font(run, Pt(9), RGBColor(0x88, 0x88, 0x88))

    # 页脚（页码）
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_prefix = footer_para.add_run('第 ')
    set_run_font(run_prefix, Pt(9), RGBColor(0x88, 0x88, 0x88))
    # 插入页码字段
    fld_char_begin = footer_para._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run_page = footer_para.add_run()
    run_page._element.append(fld_char_begin)
    instr_text = footer_para._element.makeelement(qn('w:instrText'), {})
    instr_text.text = ' PAGE '
    run_page2 = footer_para.add_run()
    run_page2._element.append(instr_text)
    fld_char_end = footer_para._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run_page3 = footer_para.add_run()
    run_page3._element.append(fld_char_end)
    run_suffix = footer_para.add_run(' 页')
    set_run_font(run_suffix, Pt(9), RGBColor(0x88, 0x88, 0x88))

    # 默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = BODY_FONT
    font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)

    # ---- 标题 ----
    add_heading1(doc, data.get('title', '文章分析报告'))
    add_spacer(doc)

    # ---- 来源信息 ----
    if data.get('source'):
        add_paragraph(doc, data['source'], Pt(10), RGBColor(0x66, 0x66, 0x66),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if data.get('sourceUrl'):
        add_paragraph(doc, f"原文链接：{data['sourceUrl']}", Pt(9),
                       RGBColor(0x88, 0x88, 0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if data.get('sourceUrls'):
        for s in data['sourceUrls']:
            url_text = f"{s.get('title', '')} - {s.get('url', '')}"
            add_paragraph(doc, url_text, Pt(9), RGBColor(0x88, 0x88, 0x88),
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_spacer(doc)

    # ---- 各章节 ----
    for sec in data.get('sections', []):
        add_heading2(doc, sec['heading'])

        sec_type = sec.get('type', '')
        if sec_type == 'bullets':
            for item in sec.get('items', []):
                add_bullet(doc, item)
        elif sec_type == 'table':
            add_table(doc, sec)
        elif sec_type == 'paragraphs':
            for sub in sec.get('subsections', []):
                add_heading3(doc, sub['subheading'])
                for para_text in sub.get('paragraphs', []):
                    add_paragraph(doc, para_text)
                add_spacer(doc)

        add_spacer(doc)

    # ---- 结尾 ----
    add_paragraph(doc, '—— 报告结束 ——', Pt(10), RGBColor(0xAA, 0xAA, 0xAA),
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


# ============================================================
# 主入口
# ============================================================

def main():
    data = load_input()
    doc = build_document(data)
    doc.save(data['outputFile'])
    print(f"✅ 文档已生成：{data['outputFile']}")


if __name__ == '__main__':
    main()
